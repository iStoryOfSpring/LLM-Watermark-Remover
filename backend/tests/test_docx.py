from pathlib import Path
from zipfile import ZipFile

from docx import Document

from backend.app.core.models import DecisionEnvelope, SentenceContext
from backend.app.document.docx import DocxAdapter
from backend.app.rewrite.runtime import ModelRuntime
from backend.app.service.rewrite_service import RewriteService


class ReplaceRuntime(ModelRuntime):
    def status(self):
        return {"state": "ready", "backend": "test"}

    def propose(self, context: SentenceContext):
        return DecisionEnvelope(
            schema_version="1.0",
            decisions=[
                {
                    "id": candidate.id,
                    "action": "replace",
                    "replacement": candidate.allowed_replacements[0],
                    "reason": "test replacement",
                }
                for candidate in context.candidates
            ],
        )


def test_docx_patches_body_text_only_and_keeps_other_parts(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("该方案")
    bold = paragraph.add_run("能够")
    bold.bold = True
    paragraph.add_run("有效提高数据处理效率。")
    document.add_heading("标题能够保留。", level=1)
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "表格能够保留。"
    document.save(source)

    with ZipFile(source) as archive:
        original_parts = {name: archive.read(name) for name in archive.namelist() if name != "word/document.xml"}

    result = RewriteService(runtime=ReplaceRuntime()).rewrite_file(source, job_dir=tmp_path / "job")
    assert result.success is True
    assert result.audit.changed == 1
    assert result.audit.protected >= 2

    output = Path(result.output_file)
    snapshot = DocxAdapter().load(output)
    assert snapshot.units[0].text == "该方案可以有效提高数据处理效率。"
    assert snapshot.units[1].text == "标题能够保留。"

    with ZipFile(output) as archive:
        for name, original_bytes in original_parts.items():
            assert archive.read(name) == original_bytes

